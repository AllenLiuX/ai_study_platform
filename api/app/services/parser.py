"""文件解析:把上传文件的字节流提取为纯文本。

支持的类型:
- application/pdf  (`pypdf`,扫描版抽不到字时上层会兜底走 vision)
- text/plain      (utf-8 解码)
- text/markdown   (utf-8 解码)
- .docx (Word 现代格式, `python-docx`)
- image/*         (Phase 4.1:占位 — 实际抽取走 services.vision_extractor)

`detect_kind` 返回 'pdf' / 'text' / 'docx' / 'doc_legacy' / 'image' / None。
'doc_legacy' 表示老 .doc 二进制 (OLE),目前不支持,上层应给出友好提示。
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass

from pypdf import PdfReader

logger = logging.getLogger(__name__)


SUPPORTED_MIME_TYPES: dict[str, str] = {
    "application/pdf": "pdf",
    "text/plain": "text",
    "text/markdown": "text",
    "text/x-markdown": "text",
    # 浏览器对 .md 经常吐空 mime,后端补一个回退
    "application/octet-stream": "binary",
    # Phase 4.2: Word 现代格式 (.docx, Office Open XML)
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    # Phase 4.2: 老 .doc 二进制(OLE),不支持但要识别出来给提示
    "application/msword": "doc_legacy",
    # Phase 4.1: 图片资料 (题目卡 / 板书 / 错题截图)
    "image/png": "image",
    "image/jpeg": "image",
    "image/jpg": "image",
    "image/webp": "image",
    "image/gif": "image",
}


SUPPORTED_EXTENSIONS: dict[str, str] = {
    "pdf": "pdf",
    "txt": "text",
    "md": "text",
    "markdown": "text",
    "docx": "docx",
    "doc": "doc_legacy",
    "png": "image",
    "jpg": "image",
    "jpeg": "image",
    "webp": "image",
    "gif": "image",
}


class EmptyPdfTextError(ValueError):
    """pypdf 抽不到任何可读文本(扫描版 / 图片型 PDF)。
    上层捕获后会触发 vision OCR 兜底。"""


@dataclass(slots=True)
class ParseResult:
    text: str
    char_count: int
    page_count: int | None
    detected_kind: str


def detect_kind(mime_type: str, filename: str) -> str | None:
    """返回 'pdf' / 'text' / 'image' / None。

    优先看 MIME,fallback 看扩展名。返回 None 表示完全不支持。
    """
    mime = (mime_type or "").lower().split(";")[0].strip()
    kind = SUPPORTED_MIME_TYPES.get(mime)
    if kind == "binary":
        kind = None
    if kind:
        return kind
    if "." in filename:
        ext = filename.rsplit(".", 1)[-1].lower()
        return SUPPORTED_EXTENSIONS.get(ext)
    return None


def parse_bytes(*, data: bytes, mime_type: str, filename: str) -> ParseResult:
    """解析文件字节流。返回纯文本与元信息。

    类型不支持抛 ValueError;PDF 抽空抛 EmptyPdfTextError;
    其他解析失败也抛 ValueError(由上层捕获写入 parse_status='failed')。

    注意:image 类型这里不做实际抽取 — vision 调用是 async,
    必须由上层 `material_processor` 直接 await `vision_extractor`。
    这个函数遇到 image kind 会抛 `NotImplementedError`,以示信号。
    """
    kind = detect_kind(mime_type, filename)
    if kind == "pdf":
        return _parse_pdf(data)
    if kind == "text":
        return _parse_text(data)
    if kind == "docx":
        return _parse_docx(data)
    if kind == "doc_legacy":
        raise ValueError(
            "暂不支持老版 .doc 二进制格式(Word 97-2003)。"
            "请在 Word/WPS 里『另存为 .docx』后再上传。"
        )
    if kind == "image":
        raise NotImplementedError(
            "image 类型必须由上层走 vision_extractor (异步),不应进入 parse_bytes"
        )
    raise ValueError(
        f"暂不支持的文件类型: mime={mime_type!r}, filename={filename!r}。"
        "目前支持 PDF / DOCX / TXT / Markdown / 图片 (PNG/JPG/WEBP/GIF)。"
    )


def _parse_pdf(data: bytes) -> ParseResult:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # PDF 损坏 / 加密
        raise ValueError(f"无法解析 PDF: {exc}") from exc

    page_texts: list[str] = []
    for idx, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            logger.warning("PDF page %d extract failed: %s", idx, exc)
            text = ""
        if text.strip():
            page_texts.append(text.strip())

    full = "\n\n".join(page_texts).strip()
    if not full:
        # 上层会捕获 EmptyPdfTextError → 走 vision_extractor 兜底
        raise EmptyPdfTextError(
            "PDF 中没有提取到可读文本,可能是扫描件或图片型 PDF;将走视觉提取兜底。"
        )
    return ParseResult(
        text=full,
        char_count=len(full),
        page_count=len(reader.pages),
        detected_kind="pdf",
    )


def _parse_docx(data: bytes) -> ParseResult:
    """解析 .docx (Office Open XML)。

    用 python-docx 抽段落 + 表格;不渲染样式,纯文本即可,后续切片不依赖排版。
    """
    try:
        from docx import Document  # 延迟导入,失败时给清晰错误
    except ImportError as exc:
        raise ValueError(
            "服务端未安装 python-docx,无法解析 .docx。请联系管理员。"
        ) from exc

    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:  # 损坏 / 加密 / 不是合法 docx
        raise ValueError(f"无法解析 .docx: {exc}") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)

    full = "\n\n".join(parts).strip()
    if not full:
        raise ValueError(".docx 中没有抽到任何可读文本,文档可能为空或只含图片。")
    return ParseResult(
        text=full,
        char_count=len(full),
        page_count=None,
        detected_kind="docx",
    )


def _parse_text(data: bytes) -> ParseResult:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            text = data.decode(encoding).strip()
            if text:
                return ParseResult(
                    text=text,
                    char_count=len(text),
                    page_count=None,
                    detected_kind="text",
                )
        except UnicodeDecodeError:
            continue
    raise ValueError("文本文件解码失败,请确认编码是 UTF-8 或 GBK。")
